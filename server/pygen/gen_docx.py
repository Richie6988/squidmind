#!/usr/bin/env python3
"""gen_docx.py <spec.json> — generate a document INHERITING a template's
styles (fonts, colors, heading formats, page setup, headers/footers).

Spec: { "template": "/abs/path.docx" | null, "output": "/abs/out.docx",
        "title": str?, "markdown": str, "asset_root": "/abs/aquarium"?,
        "charts": [{"type": "bar"|"line"|"pie", "categories": [str],
                    "series": [{"name": str, "values": [num]}]}]? }

Markdown extensions:
  ![alt](path)      → embedded picture (path resolved against asset_root,
                      confined to it), fitted to the page width
  | a | b | tables  → real Word table using the template's table style
  {{chart:N}}       → charts[N] rendered to PNG via matplotlib and embedded
                      (matplotlib needed only when charts are used)

Strategy:
  1. Open the template with python-docx: styles, theme, sectPr (margins,
     header/footer with logos) come along.
  2. Clear the template's BODY content (paragraphs + tables) while keeping
     the final sectPr — content was reference, chrome stays.
  3. Append the markdown using the template's OWN named styles:
     'Title', 'Heading 1..3', 'List Bullet', 'List Number', 'Normal'.
     Missing style → graceful fallback to Normal. Inline **bold** and
     *italic* become runs.
Result JSON on stdout.
"""
import json, os, re, sys, tempfile

def fail(msg):
    print(json.dumps({"ok": False, "error": str(msg)})); sys.exit(0)

try:
    import docx
    from docx import Document
except Exception as e:
    fail(f"python-docx not available: {e}")

ASSET_ROOT = None

def resolve_asset(p):
    if not p: raise ValueError("empty image path")
    cand = p if os.path.isabs(p) else os.path.join(ASSET_ROOT or "", p)
    real = os.path.realpath(cand)
    if ASSET_ROOT and not real.startswith(os.path.realpath(ASSET_ROOT) + os.sep):
        raise ValueError(f"image path escapes the aquarium: {p}")
    if not os.path.isfile(real):
        raise ValueError(f"image not found: {p}")
    return real

def usable_width(doc):
    s = doc.sections[-1]
    return s.page_width - s.left_margin - s.right_margin

def table_style(doc):
    for name in ("Table Grid", "Light Grid", "Light List"):
        try:
            doc.styles[name]; return name
        except KeyError: continue
    return None

def render_chart_png(spec):
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except Exception as e:
        raise ValueError(f"charts need matplotlib in the IAQUA venv: {e}")
    cats   = [str(c) for c in (spec.get("categories") or [])]
    series = spec.get("series") or []
    fig, ax = plt.subplots(figsize=(6.4, 3.6), dpi=140)
    kind = str(spec.get("type", "bar")).lower()
    if kind == "pie" and series:
        ax.pie([float(v) for v in series[0].get("values") or []], labels=cats, autopct="%1.0f%%")
    elif kind == "line":
        for s in series:
            ax.plot(cats, [float(v) for v in s.get("values") or []], marker="o", label=str(s.get("name", "")))
    else:
        import numpy as _np
        x = _np.arange(len(cats)); n = max(1, len(series)); w = 0.8 / n
        for i, s in enumerate(series):
            ax.bar(x + i * w - 0.4 + w / 2, [float(v) for v in s.get("values") or []], w, label=str(s.get("name", "")))
        ax.set_xticks(x); ax.set_xticklabels(cats)
    if kind != "pie" and len(series) > 1: ax.legend()
    fig.tight_layout()
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    fig.savefig(tmp.name); plt.close(fig)
    return tmp.name

def flush_table(doc, tbl_lines):
    """tbl_lines: raw markdown | rows | — build a real Word table."""
    rows = []
    for ln in tbl_lines:
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        if cells and all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
            continue  # separator row
        rows.append(cells)
    if not rows: return
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncols)
    st = table_style(doc)
    if st: t.style = st
    for ri, r in enumerate(rows):
        for ci in range(ncols):
            cell = t.cell(ri, ci)
            cell.text = r[ci] if ci < len(r) else ""
            if ri == 0:
                for p in cell.paragraphs:
                    for run in p.runs: run.bold = True

def clear_body(doc):
    body = doc.element.body
    # Keep the trailing sectPr (page setup / headers-footers), drop the rest.
    for child in list(body):
        if child.tag.endswith('}sectPr'):
            continue
        body.remove(child)

def style_or(doc, name, fallback=None):
    try:
        doc.styles[name]
        return name
    except KeyError:
        return fallback

INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")

def add_runs(par, text):
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            par.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            par.add_run(tok[2:-2]).bold = True
        else:
            par.add_run(tok[1:-1]).italic = True
        pos = m.end()
    if pos < len(text):
        par.add_run(text[pos:])

def main():
    global ASSET_ROOT
    spec = json.load(open(sys.argv[1], encoding="utf-8"))
    ASSET_ROOT = spec.get("asset_root")
    doc = Document(spec["template"]) if spec.get("template") else Document()
    if spec.get("template"): clear_body(doc)
    charts = spec.get("charts") or []

    S = {
        "title": style_or(doc, "Title"),
        "h1": style_or(doc, "Heading 1"),
        "h2": style_or(doc, "Heading 2"),
        "h3": style_or(doc, "Heading 3"),
        "bullet": style_or(doc, "List Bullet"),
        "number": style_or(doc, "List Number"),
    }

    if spec.get("title"):
        p = doc.add_paragraph()
        if S["title"]: p.style = S["title"]
        add_runs(p, spec["title"])

    para_buf = []
    def flush_para():
        nonlocal para_buf
        if para_buf:
            p = doc.add_paragraph()
            add_runs(p, " ".join(para_buf))
            para_buf = []

    tbl_buf = []
    def flush_tbl():
        nonlocal tbl_buf
        if tbl_buf:
            flush_para()
            flush_table(doc, tbl_buf)
            tbl_buf = []

    IMG_RE   = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")
    CHART_RE = re.compile(r"^\{\{chart:(\d+)\}\}\s*$")

    for raw in spec.get("markdown", "").splitlines():
        line = raw.rstrip()
        # Markdown table rows accumulate, anything else flushes them
        if line.strip().startswith("|") and line.strip().endswith("|"):
            tbl_buf.append(line); continue
        flush_tbl()
        m = IMG_RE.match(line.strip())
        if m:
            flush_para()
            real = resolve_asset(m.group(2).strip())
            doc.add_picture(real, width=usable_width(doc))
            if m.group(1).strip():
                cap = doc.add_paragraph()
                cap.add_run(m.group(1).strip()).italic = True
            continue
        m = CHART_RE.match(line.strip())
        if m:
            flush_para()
            idx = int(m.group(1))
            if idx >= len(charts): raise ValueError(f"{{{{chart:{idx}}}}} but only {len(charts)} chart spec(s) given")
            png = render_chart_png(charts[idx])
            doc.add_picture(png, width=usable_width(doc))
            try: os.unlink(png)
            except OSError: pass
            continue
        if not line.strip():
            flush_para(); continue
        m = re.match(r"^(#{1,3})\s+(.*)$", line)
        if m:
            flush_para()
            lvl = len(m.group(1))
            p = doc.add_paragraph()
            st = S["h1"] if lvl == 1 else S["h2"] if lvl == 2 else S["h3"]
            if st: p.style = st
            add_runs(p, m.group(2))
            continue
        m = re.match(r"^\s*[-*]\s+(.*)$", line)
        if m:
            flush_para()
            p = doc.add_paragraph(style=S["bullet"]) if S["bullet"] else doc.add_paragraph()
            add_runs(p, m.group(1))
            continue
        m = re.match(r"^\s*\d+[.)]\s+(.*)$", line)
        if m:
            flush_para()
            p = doc.add_paragraph(style=S["number"]) if S["number"] else doc.add_paragraph()
            add_runs(p, m.group(1))
            continue
        para_buf.append(line.strip())
    flush_tbl()
    flush_para()

    doc.save(spec["output"])
    print(json.dumps({"ok": True, "outputPath": spec["output"], "engine": "python-docx (template-inherited)" if spec.get("template") else "python-docx (blank)"}))

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        fail(e)
